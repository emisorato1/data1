from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType

from .base import DocumentLoader
from .models import LoadedDocument


class DOCXLoader(DocumentLoader):
    """Cargador de documentos DOCX usando python-docx."""

    def load(self, file_path: Path) -> LoadedDocument:
        """Extrae texto, metadata y tablas de un archivo Word."""
        try:
            doc = Document(str(file_path))
            full_text = []
            tables_count = 0

            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            for table in doc.tables:
                tables_count += 1
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append("| " + " | ".join(row_data) + " |")

                if table_data:
                    num_cols = len(table.rows[0].cells)
                    separator = "| " + " | ".join(["---"] * num_cols) + " |"
                    table_md = [table_data[0], separator, *table_data[1:]]
                    full_text.append("\n[TABLA]\n" + "\n".join(table_md) + "\n[/TABLA]\n")

            metadata = self._extract_metadata(doc, file_path)

            pages_estimate = self._estimate_pages(doc)

            return LoadedDocument(
                text="\n".join(full_text), metadata=metadata, pages=pages_estimate, tables_count=tables_count
            )
        except Exception as e:
            raise Exception(f"Error cargando DOCX {file_path}: {e}") from e

    def _extract_metadata(self, doc: DocumentType, file_path: Path) -> dict[str, Any]:
        cp = doc.core_properties
        return {
            "source": str(file_path.name),
            "title": cp.title or file_path.stem,
            "author": cp.author or "Unknown",
            "created": str(cp.created) if cp.created else None,
            "modified": str(cp.modified) if cp.modified else None,
            "file_size": file_path.stat().st_size,
        }

    def _estimate_pages(self, doc: DocumentType) -> int:
        """Estima el número de páginas basado en word count (~250 palabras/página)."""
        word_count = sum(len(p.text.split()) for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    word_count += len(cell.text.split())
        return max(1, (word_count + 249) // 250)
