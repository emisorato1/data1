from docx import Document as DocxDocument
from models.document import Document


def load_docx(path: str) -> Document:
    """Carga un archivo DOCX y devuelve un Document."""
    doc = DocxDocument(path)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)

    metadata = {
        "source": path,
        "file_type": "docx",
        "num_paragraphs": len(paragraphs),
    }

    return Document(content=text, metadata=metadata)
